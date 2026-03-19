"""DOCX template filler for Polish employment contracts.

Handles {{PLACEHOLDER}} replacement in Word documents,
including across split runs within paragraphs and table cells.
"""
import io
import re
import logging
from copy import deepcopy

from docx import Document

_logger = logging.getLogger(__name__)

_PLACEHOLDER_RE = re.compile(r"\{\{(\w+)\}\}")


def fill_docx_template(template_bytes: bytes, placeholders: dict) -> bytes:
    """Fill a DOCX template with placeholder values.

    Args:
        template_bytes: Raw bytes of the DOCX template file.
        placeholders: Dict mapping placeholder names (without braces) to values.
                     e.g. {"IMIE_I_NAZWISKO": "Jan Kowalski", "PESEL": "12345678901"}

    Returns:
        Filled DOCX as bytes.
    """
    doc = Document(io.BytesIO(template_bytes))

    # Process all paragraphs in the document body
    for para in doc.paragraphs:
        _replace_in_paragraph(para, placeholders)

    # Process all tables (main content is in tables for bilingual layout)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _replace_in_paragraph(para, placeholders)

    # Process headers and footers
    for section in doc.sections:
        for header in (section.header, section.first_page_header, section.even_page_header):
            if header:
                for para in header.paragraphs:
                    _replace_in_paragraph(para, placeholders)
                for table in header.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for para in cell.paragraphs:
                                _replace_in_paragraph(para, placeholders)
        for footer in (section.footer, section.first_page_footer, section.even_page_footer):
            if footer:
                for para in footer.paragraphs:
                    _replace_in_paragraph(para, placeholders)
                for table in footer.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for para in cell.paragraphs:
                                _replace_in_paragraph(para, placeholders)

    # Save to bytes
    output = io.BytesIO()
    doc.save(output)
    return output.getvalue()


def _replace_in_paragraph(paragraph, placeholders: dict):
    """Replace {{PLACEHOLDER}} patterns in a paragraph, handling split runs."""
    # First check if there are any placeholders in the full text
    full_text = "".join(run.text for run in paragraph.runs)
    if "{{" not in full_text:
        return

    # Check if any replacements would actually happen
    new_text = _apply_replacements(full_text, placeholders)
    if new_text == full_text:
        return

    # Strategy: merge all runs into the first run, clear the rest.
    # This preserves the formatting of the first run.
    # For bilingual contracts this is fine since each cell has consistent formatting.
    if not paragraph.runs:
        return

    # Preserve formatting of the first non-empty run
    fmt_run = None
    for run in paragraph.runs:
        if run.text.strip():
            fmt_run = run
            break
    if fmt_run is None:
        fmt_run = paragraph.runs[0]

    # Handle multiline replacements (e.g., scope of work)
    if "\n" in new_text:
        lines = new_text.split("\n")
        # Set first line in the first run
        fmt_run.text = lines[0]
        # Clear remaining runs
        for run in paragraph.runs:
            if run is not fmt_run:
                run.text = ""
        # Add line breaks for remaining lines
        from docx.oxml.ns import qn
        for line in lines[1:]:
            br = fmt_run._element.makeelement(qn("w:br"), {})
            fmt_run._element.append(br)
            # Add text after break using a new run-like approach
            t = fmt_run._element.makeelement(qn("w:t"), {})
            t.text = line
            t.set(qn("xml:space"), "preserve")
            fmt_run._element.append(t)
    else:
        # Simple case: set text on first run, clear others
        fmt_run.text = new_text
        for run in paragraph.runs:
            if run is not fmt_run:
                run.text = ""


def _apply_replacements(text: str, placeholders: dict) -> str:
    """Replace all {{KEY}} patterns in text with values from placeholders dict."""
    def replacer(match):
        key = match.group(1)
        if key in placeholders:
            val = placeholders[key]
            return str(val) if val else ""
        # Leave unreplaced if key not in dict (template may have conditional fields)
        return match.group(0)

    return _PLACEHOLDER_RE.sub(replacer, text)


def get_unfilled_placeholders(template_bytes: bytes) -> list:
    """Extract all unique placeholder names from a DOCX template.

    Useful for validation — check which placeholders the template expects.
    """
    doc = Document(io.BytesIO(template_bytes))
    found = set()

    # Check paragraphs
    for para in doc.paragraphs:
        full_text = "".join(run.text for run in para.runs)
        found.update(_PLACEHOLDER_RE.findall(full_text))

    # Check tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    full_text = "".join(run.text for run in para.runs)
                    found.update(_PLACEHOLDER_RE.findall(full_text))

    return sorted(found)
